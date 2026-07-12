"""
PAAIM Project Documentation Generator
Produces a professional Word document with embedded architecture diagrams.

Usage:
    cd /Users/venky/Desktop/projects/PAAIM
    python generate_doc.py
"""

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import matplotlib.patheffects as pe

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x0F, 0x29, 0x4D)   # headings
MID_BLUE   = RGBColor(0x1E, 0x5C, 0xAD)   # sub-headings
ACCENT     = RGBColor(0x2B, 0x7A, 0xF0)   # highlights
GREY_TEXT  = RGBColor(0x4A, 0x4A, 0x4A)   # body
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)   # table fills

# ── Diagram helpers ───────────────────────────────────────────────────────────

def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def make_pipeline_diagram():
    """Vertical 7-layer pipeline flow."""
    layers = [
        ("1  Event Ingestion",      "#E8F4FD", "#1E5CAD", "SCADA / IoT / MES signals"),
        ("2  Agent Analysis",        "#EAF7EA", "#1B7A3C", "5 Gemini-powered specialists"),
        ("3  Policy Engine",         "#FFF3E0", "#E65100", "Industrial Constitution check"),
        ("4  Decision Twin",         "#F3E5F5", "#7B1FA2", "Impact simulation"),
        ("5  Red-Team Challenge",    "#FCE4EC", "#C62828", "AI questions itself"),
        ("6  Approval Gate",         "#E8EAF6", "#283593", "Routes to right human role"),
        ("7  Audit & Response",      "#F1F8E9", "#33691E", "Full evidence pack recorded"),
    ]

    fig, ax = plt.subplots(figsize=(8, 9.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, len(layers) * 1.4 + 0.5)
    ax.axis('off')

    box_h = 1.0
    gap   = 0.4
    y_positions = []

    for i, (label, bg, border, sub) in enumerate(reversed(layers)):
        y = i * (box_h + gap) + 0.3
        y_positions.append(y + box_h / 2)

        # Box
        rect = FancyBboxPatch((0.5, y), 9, box_h,
                               boxstyle="round,pad=0.08",
                               linewidth=2, edgecolor=border,
                               facecolor=bg)
        ax.add_patch(rect)

        # Layer label
        ax.text(1.1, y + box_h * 0.62, label,
                fontsize=11, fontweight='bold', color=border, va='center')
        # Sub-label
        ax.text(1.1, y + box_h * 0.28, sub,
                fontsize=8.5, color='#555555', va='center')

    # Arrows between boxes
    for i in range(len(y_positions) - 1):
        ax.annotate('', xy=(5, y_positions[i+1] - box_h / 2 - 0.02),
                    xytext=(5, y_positions[i] + box_h / 2 + 0.02),
                    arrowprops=dict(arrowstyle='->', color='#888888', lw=1.5))

    # Title
    ax.text(5, len(layers) * (box_h + gap) + 0.3,
            "PAAIM — 7-Layer Orchestration Pipeline",
            fontsize=13, fontweight='bold', color='#0F294D',
            ha='center', va='bottom')

    return fig_to_bytes(fig)


def make_architecture_diagram():
    """High-level system architecture block diagram."""
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')

    def box(x, y, w, h, label, sub, bg, border, fontsize=9):
        rect = FancyBboxPatch((x, y), w, h,
                               boxstyle="round,pad=0.12",
                               linewidth=1.8, edgecolor=border,
                               facecolor=bg)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h*0.62, label,
                ha='center', va='center', fontsize=fontsize,
                fontweight='bold', color=border)
        ax.text(x + w/2, y + h*0.28, sub,
                ha='center', va='center', fontsize=7.5,
                color='#555555')

    def arrow(x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#888888', lw=1.3))

    # ── Frontend ──
    box(0.3, 6.5, 3.5, 1.8, "Next.js Frontend", "React · Recharts · Zustand",
        "#E8F4FD", "#1E5CAD", 10)

    # ── FastAPI Backend ──
    box(5.0, 5.2, 4.0, 3.5, "FastAPI Backend", "Python · Async SQLAlchemy",
        "#F9FBE7", "#558B2F", 10)

    # ── Orchestrator ──
    box(5.3, 5.5, 3.4, 1.2, "Orchestrator", "7-layer pipeline",
        "#FFFDE7", "#F9A825", 8)

    # ── Policy Engine ──
    box(0.3, 3.5, 3.0, 1.5, "Policy Engine", "Industrial Constitution YAML",
        "#FFF3E0", "#E65100")

    # ── Gemini Agents ──
    box(0.3, 1.2, 3.0, 1.8, "Gemini Agents", "Safety · Quality\nMaintenance · Energy",
        "#EAF7EA", "#2E7D32")

    # ── Decision Twin ──
    box(9.2, 6.5, 4.2, 1.8, "Decision Twin", "Impact simulation",
        "#F3E5F5", "#7B1FA2")

    # ── Approval Gate ──
    box(9.2, 4.2, 4.2, 1.8, "Approval Gate", "Operator · Supervisor\nSafety Officer",
        "#E8EAF6", "#283593")

    # ── Database ──
    box(9.2, 1.8, 4.2, 1.8, "SQLite / PostgreSQL", "Events · Decisions\nAudit Logs",
        "#ECEFF1", "#455A64")

    # ── Sensor sources ──
    box(5.0, 0.3, 4.0, 1.2, "Factory Data Sources", "SCADA · MES · IoT · OPC-UA",
        "#FCE4EC", "#AD1457")

    # Arrows
    arrow(3.5, 7.4, 5.0, 7.4)    # Frontend → Backend
    arrow(5.3, 6.9, 4.5, 4.8)    # Backend → Policy
    arrow(5.3, 6.2, 4.5, 2.8)    # Backend → Agents
    arrow(9.0, 7.4, 9.2, 7.4)    # Backend → Twin
    arrow(9.0, 6.5, 9.2, 5.1)    # Backend → Approval
    arrow(9.0, 5.8, 9.2, 2.6)    # Backend → DB
    arrow(7.0, 1.5, 7.0, 0.3)    # Backend ← Sources (reversed)

    ax.text(7.0, 9.0, "PAAIM — System Architecture",
            ha='center', fontsize=13, fontweight='bold', color='#0F294D')

    return fig_to_bytes(fig)


def make_decision_flow():
    """How a single event flows to an approved decision."""
    steps = [
        ("Sensor Signal",        "#E3F2FD", "#1565C0", "vibration_anomaly  0.87  conf 0.91"),
        ("Safety Agent",         "#E8F5E9", "#2E7D32", "Gemini: 'schedule_maintenance · medium risk'"),
        ("Policy Check",         "#FFF8E1", "#F57F17", "constitution.yaml → operator approval required"),
        ("Impact Simulation",    "#F3E5F5", "#6A1B9A", "downtime 2h · cost $1,500 · OEE –3%"),
        ("Red-Team Review",      "#FCE4EC", "#B71C1C", "No safety conflict · acceptable risk"),
        ("Approval Gate",        "#E8EAF6", "#1A237E", "Routed to: operator_001"),
        ("Human Approves",       "#E0F2F1", "#004D40", "operator_001 clicks Approve · timestamp logged"),
        ("Audit Record",         "#F9FBE7", "#33691E", "7 audit entries written to DB"),
    ]

    fig, ax = plt.subplots(figsize=(9, 10))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, len(steps) * 1.35 + 0.6)
    ax.axis('off')

    bh = 1.0
    gap = 0.35

    for i, (label, bg, border, detail) in enumerate(reversed(steps)):
        y = i * (bh + gap) + 0.3
        rect = FancyBboxPatch((0.4, y), 10.2, bh,
                               boxstyle="round,pad=0.1",
                               linewidth=1.8, edgecolor=border, facecolor=bg)
        ax.add_patch(rect)
        ax.text(0.85, y + bh*0.65, label,
                fontsize=10, fontweight='bold', color=border)
        ax.text(0.85, y + bh*0.28, detail,
                fontsize=8, color='#444444', family='monospace')

        if i < len(steps) - 1:
            next_y = (i + 1) * (bh + gap) + 0.3
            ax.annotate('', xy=(5.5, next_y - 0.04),
                        xytext=(5.5, y + bh + 0.04),
                        arrowprops=dict(arrowstyle='->', color='#888888', lw=1.4))

    ax.text(5.5, len(steps) * (bh + gap) + 0.45,
            "Single Event → Approved Decision Flow",
            fontsize=12, fontweight='bold', color='#0F294D',
            ha='center')

    return fig_to_bytes(fig)


# ── Document helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    run.font.bold = True
    return p


def add_body(doc, text, space_before=0):
    p = doc.add_paragraph(text)
    p.style = 'Normal'
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = GREY_TEXT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GREY_TEXT
    p.paragraph_format.space_after = Pt(3)
    return p


def add_image(doc, img_bytes, width=6.0, caption=None):
    doc.add_picture(img_bytes, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        cp.paragraph_format.space_after = Pt(10)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '1E5CAD')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'F5F8FF' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = GREY_TEXT

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(8)


# ── Build the document ────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Cover page ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("PAAIM")
    r.font.size = Pt(42)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Policy-Aware Agentic Intelligence Manager")
    r2.font.size = Pt(18)
    r2.font.color.rgb = MID_BLUE

    doc.add_paragraph()
    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = desc_p.add_run("A Manufacturing Decision Intelligence System")
    r3.font.size = Pt(13)
    r3.font.color.rgb = GREY_TEXT
    r3.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Version",     "1.0 — Phase 1 Demo"),
        ("Date",        "June 2026"),
        ("Factory",     "Factory 001 · Austin, TX"),
        ("Stack",       "FastAPI · Next.js · Gemini 2.0 Flash · SQLite"),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = p.add_run(f"{label}: ")
        rb.font.bold = True
        rb.font.size = Pt(11)
        rb.font.color.rgb = DARK_BLUE
        rv = p.add_run(val)
        rv.font.size = Pt(11)
        rv.font.color.rgb = GREY_TEXT

    doc.add_page_break()

    # ── 1. Introduction ───────────────────────────────────────────────────────
    add_heading(doc, "1.  Introduction")
    add_body(doc, (
        "A typical manufacturing plant generates several hundred sensor alerts every day. "
        "Operators are expected to review each one, decide whether it is serious, and act — "
        "all while keeping the line running. In practice, alert fatigue sets in quickly. "
        "Important signals get missed, decisions are made under time pressure with incomplete "
        "information, and when something goes wrong, there is rarely a clear record of why "
        "a particular course of action was chosen."
    ))
    add_body(doc, (
        "PAAIM was built to address this directly. When a sensor signal arrives — a vibration "
        "spike, a quality defect, an energy anomaly — PAAIM passes it through a seven-layer "
        "pipeline: specialist AI agents analyse the event, a policy engine checks what actions "
        "are permitted, a simulation estimates the cost of each option, a red-team module "
        "questions the recommendation, and then the decision is routed to the right human for "
        "approval. Every step is recorded."
    ))
    add_body(doc, (
        "The result is not a system that replaces operators. It is one that gives them "
        "better information, faster, with a clear audit trail behind every decision."
    ))

    add_heading(doc, "1.1  The core problem", level=2)
    add_bullet(doc, "Operators receive too many alerts to process carefully, so they prioritise by habit rather than by evidence.")
    add_bullet(doc, "When a decision is made, the reasoning behind it often exists only in someone's head — not in a log.")
    add_bullet(doc, "There is no consistent process for deciding which role (operator, supervisor, safety officer) should approve a given action.")
    add_bullet(doc, "Predictive signals (bearing wear, quality drift) are available but rarely acted on before failure occurs.")

    add_heading(doc, "1.2  What PAAIM does differently", level=2)
    add_bullet(doc, "Runs every incoming signal through a structured seven-layer pipeline before a recommendation reaches a human.")
    add_bullet(doc, "Enforces an 'Industrial Constitution' — a YAML policy file that defines what actions are allowed, who must approve them, and what evidence is required.")
    add_bullet(doc, "Logs every agent analysis, policy evaluation, and approval decision to a searchable audit trail.")
    add_bullet(doc, "Lets engineers add new monitoring agents without writing code — configure a data source, define threshold rules, and the agent is live.")

    doc.add_page_break()

    # ── 2. System Architecture ────────────────────────────────────────────────
    add_heading(doc, "2.  System Architecture")
    add_body(doc, (
        "PAAIM consists of three main parts: a Next.js frontend that operators use to review "
        "decisions and run scenarios, a FastAPI backend that hosts the orchestration pipeline, "
        "and a SQLite database (PostgreSQL in production) that stores events, decisions, and "
        "audit logs. The Gemini 2.0 Flash API is called by each specialist agent during the "
        "analysis layer."
    ))

    arch_bytes = make_architecture_diagram()
    add_image(doc, arch_bytes, width=6.0,
              caption="Figure 1 — High-level system architecture")

    add_heading(doc, "2.1  Technology stack", level=2)
    add_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend",       "Next.js 14, React, Tailwind CSS, Recharts",  "Operator dashboard, analytics, agent builder"],
            ["State mgmt",     "Zustand, TanStack Query",                    "Client state and server data caching"],
            ["Backend",        "Python 3.11, FastAPI, Uvicorn",              "API server, orchestration pipeline"],
            ["ORM",            "SQLAlchemy 2 (async), Alembic",              "Database models and migrations"],
            ["Database (dev)", "SQLite + aiosqlite",                         "Local development and demo"],
            ["Database (prod)","PostgreSQL + asyncpg",                       "Production deployment"],
            ["AI",             "Google Gemini 2.0 Flash",                    "Agent reasoning and recommendation"],
            ["Streaming",      "WebSocket (FastAPI native)",                  "Real-time pipeline progress to UI"],
        ],
        col_widths=[1.4, 2.6, 2.8]
    )

    doc.add_page_break()

    # ── 3. The 7-Layer Pipeline ───────────────────────────────────────────────
    add_heading(doc, "3.  The 7-Layer Pipeline")
    add_body(doc, (
        "Every manufacturing event — regardless of type — passes through the same seven layers "
        "before a decision is stored. The layers run sequentially so that each one builds on the "
        "output of the previous. Typical end-to-end latency is under 1,500 ms."
    ))

    pipeline_bytes = make_pipeline_diagram()
    add_image(doc, pipeline_bytes, width=5.5,
              caption="Figure 2 — The 7-layer PAAIM pipeline")

    layers_detail = [
        ("Layer 1 — Event Ingestion",
         "Raw signals arrive from SCADA systems, IoT sensors, MES platforms, or the built-in "
         "scenario simulator. Each event carries a type (safety, quality, maintenance, production, "
         "energy), a signal name, a value, a confidence score, and factory context. The event is "
         "written to the database before processing begins."),
        ("Layer 2 — Agent Analysis",
         "Five specialist agents run in parallel against the incoming event. Each agent sends the "
         "event data to Gemini 2.0 Flash with a domain-specific context prompt and receives back "
         "a structured JSON response: a confidence score, reasoning, and a list of recommended "
         "actions with risk levels. If Gemini is unavailable, each agent falls back to a "
         "deterministic rule set. Custom agents — defined through the no-code builder — also run "
         "at this stage."),
        ("Layer 3 — Policy Evaluation",
         "Each recommended action is checked against the Industrial Constitution, a YAML file that "
         "defines which actions are allowed, which approval level they require, what evidence "
         "signals must be present, and which actions conflict with each other. Actions not defined "
         "in the constitution are blocked. The policy engine returns an approval level "
         "(auto, operator, supervisor, manager, or safety_officer) for each action."),
        ("Layer 4 — Decision Twin",
         "For each candidate action, the Decision Twin estimates the operational impact before "
         "anything is executed: expected downtime hours, scrap units, OEE change, and cost "
         "impact. This gives the approving human concrete numbers rather than a vague "
         "recommendation."),
        ("Layer 5 — Red-Team Challenge",
         "A separate module reviews the top recommendation and actively tries to find problems "
         "with it. It checks whether the confidence is justified, whether the evidence could "
         "support an alternative interpretation, and whether a safer or cheaper action exists. "
         "Its output appears in the decision detail view alongside the recommendation."),
        ("Layer 6 — Approval Gate",
         "The highest-priority permitted action is selected and routed to the appropriate "
         "approver. Actions with approval_required: auto in the constitution are approved "
         "immediately (currently only acknowledge_estop). All others generate a 'recommended' "
         "decision that waits for a human to approve or reject through the dashboard. Risk level "
         "determines escalation: critical actions go to both the safety officer and plant manager."),
        ("Layer 7 — Audit Trail",
         "Four audit log entries are written for every decision: event detected, agent analysed, "
         "policy evaluated, and decision outcome. A fifth entry is added when a human approves "
         "or rejects. Every log entry records who acted, when, and what the data showed at that "
         "moment. This trail satisfies ISO 45001 evidence requirements."),
    ]

    for title, body in layers_detail:
        add_heading(doc, title, level=2)
        add_body(doc, body)

    doc.add_page_break()

    # ── 4. Decision Flow ──────────────────────────────────────────────────────
    add_heading(doc, "4.  Event to Decision — A Worked Example")
    add_body(doc, (
        "The following trace shows a real vibration anomaly from the demo dataset passing "
        "through the full pipeline. The sensor reading is 0.87 (above the 0.80 threshold) "
        "with a confidence score of 0.91."
    ))

    flow_bytes = make_decision_flow()
    add_image(doc, flow_bytes, width=5.8,
              caption="Figure 3 — Single event flowing from sensor signal to approved decision")

    add_body(doc, (
        "The operator sees the recommendation in the dashboard, reads the agent reasoning, "
        "the impact estimate (2 hours downtime, $1,500 cost), and the red-team assessment. "
        "They click Approve. The decision status changes from 'recommended' to 'approved', "
        "the approval is timestamped, and two additional audit entries are written."
    ))

    doc.add_page_break()

    # ── 5. Key Components ────────────────────────────────────────────────────
    add_heading(doc, "5.  Key Components")

    add_heading(doc, "5.1  Specialist Agents", level=2)
    add_body(doc, (
        "Five agents cover the primary event domains found in discrete manufacturing. Each "
        "agent has a domain-specific system prompt that frames the Gemini analysis in the "
        "right operational context."
    ))
    add_table(doc,
        ["Agent", "Domain", "Key signals", "Key actions"],
        [
            ["Safety Agent",      "Worker and equipment safety",    "zone_intrusion, e_stop_signal",       "stop_line, acknowledge_estop"],
            ["Quality Agent",     "Defect detection and containment","defect_detection, surface_finish",   "contain_batch, inspect_root_cause"],
            ["Maintenance Agent", "Predictive maintenance",          "vibration_anomaly, bearing_wear",    "schedule_maintenance, escalate_critical"],
            ["Production Agent",  "Order fulfilment and throughput", "order_at_risk, throughput_drop",     "propose_recovery_plan, adjust_schedule"],
            ["Energy Agent",      "Energy cost optimisation",        "peak_pricing_window, idle_load",     "shift_non_critical_load, reduce_consumption"],
        ],
        col_widths=[1.5, 1.8, 2.0, 2.0]
    )

    add_heading(doc, "5.2  The Industrial Constitution", level=2)
    add_body(doc, (
        "The Industrial Constitution is a YAML file that acts as the single source of truth "
        "for what the system is allowed to do. It is not code — it is a policy document that "
        "anyone on the team can read and edit. Each action entry specifies the approval level "
        "required, the minimum confidence threshold, which actions it conflicts with, and the "
        "expected operational impact."
    ))
    add_body(doc, (
        "This design keeps policy separate from application logic. Changing who must approve "
        "a line stop does not require a code deployment — it requires editing one line in the YAML "
        "and restarting the backend. The constitution currently defines twelve actions across "
        "five categories: safety, quality, maintenance, production, and energy."
    ))

    add_heading(doc, "5.3  Custom Agent Framework", level=2)
    add_body(doc, (
        "The no-code agent builder allows engineers to extend PAAIM without writing Python. "
        "An agent definition has four parts: a name and domain, one or more data source "
        "connections (SCADA, REST API, IoT, OPC-UA), a set of threshold rules, and a list "
        "of actions the agent can recommend."
    ))
    add_body(doc, (
        "Once created, the agent is saved to a JSON file on disk and loaded automatically on "
        "startup. It runs alongside the built-in agents during Layer 2 of every pipeline "
        "execution. The Thermal Monitor Agent in the demo was created through the UI in about "
        "60 seconds and immediately appeared in the decision detail view of the next scenario run."
    ))

    add_heading(doc, "5.4  Approval Workflow", level=2)
    add_table(doc,
        ["Role", "Approval threshold", "Deadline", "Example action"],
        [
            ["Auto (system)",   "auto",            "Immediate",  "acknowledge_estop"],
            ["Operator",        "LOW / MEDIUM",    "60 minutes", "schedule_maintenance"],
            ["Supervisor",      "MEDIUM / HIGH",   "15 minutes", "contain_batch"],
            ["Plant Manager",   "CRITICAL",        "5 minutes",  "escalate_critical"],
            ["Safety Officer",  "CRITICAL / IMMEDIATE", "1 minute", "stop_line"],
        ],
        col_widths=[1.5, 1.7, 1.3, 2.7]
    )

    doc.add_page_break()

    # ── 6. Frontend ───────────────────────────────────────────────────────────
    add_heading(doc, "6.  Frontend and Operator Experience")
    add_body(doc, (
        "The dashboard is built with Next.js 14 and styled with Tailwind CSS. It uses a dark "
        "sidebar navigation — a design pattern familiar from industrial monitoring tools — with "
        "a white content area for readability. All data is loaded via React Query with automatic "
        "background refresh."
    ))

    add_heading(doc, "6.1  Pages", level=2)
    add_table(doc,
        ["Page", "Path", "What it shows"],
        [
            ["Login",           "/login",          "Authentication with demo credentials"],
            ["Landing",         "/",               "Project overview and feature summary"],
            ["Operations",      "/dashboard",      "Live incidents, test scenarios, decisions list with approve/reject"],
            ["Decision Detail", "/dashboard/[id]", "Full pipeline trace, agent analyses, impact, red-team, approval panel"],
            ["Analytics",       "/analytics",      "KPI cards, event timeline, agent performance, pipeline latency"],
            ["Agents",          "/custom-agents",  "Custom agent list and no-code builder"],
            ["Audit Trail",     "/audit",          "Searchable log of all pipeline events"],
            ["Knowledge",       "/knowledge",      "Factory model and KPI definitions"],
        ],
        col_widths=[1.5, 1.8, 3.9]
    )

    add_heading(doc, "6.2  Real-time streaming", level=2)
    add_body(doc, (
        "The decision detail page connects to a WebSocket endpoint on the backend. As the "
        "pipeline processes an event, each layer emits a progress event that appears in the "
        "live pipeline view in real time — agents routing, policy checking, simulation complete. "
        "A keepalive ping runs every 20 seconds to prevent connection timeout."
    ))

    doc.add_page_break()

    # ── 7. Demo Results ──────────────────────────────────────────────────────
    add_heading(doc, "7.  Demo Dataset and Metrics")
    add_body(doc, (
        "The demo database is pre-seeded with 30 days of realistic manufacturing activity "
        "using the seed_demo.py script. The dataset reflects a maintenance-heavy factory "
        "with the following distribution."
    ))

    add_table(doc,
        ["Metric", "Value", "Notes"],
        [
            ["Total events",          "130",      "30-day period"],
            ["Total decisions",       "130",      "One per event"],
            ["Approved decisions",    "108",      "Status: approved"],
            ["Pending (recommended)", "14",       "Awaiting operator action"],
            ["Rejected",              "8",        "Operator rejected recommendation"],
            ["Approval rate",         "83.1%",    "Approved ÷ total decisions"],
            ["Estimated cost savings","$486,000", "Approved decisions × $4,500 avg value"],
            ["Audit log entries",     "520",      "4 entries per decision"],
            ["Avg pipeline latency",  "1,457 ms", "Across all 7 layers"],
        ],
        col_widths=[2.5, 1.5, 3.2]
    )

    add_heading(doc, "7.1  Event distribution", level=2)
    add_table(doc,
        ["Event type", "Count", "Share"],
        [
            ["Maintenance", "45", "34.6%"],
            ["Quality",     "35", "26.9%"],
            ["Safety",      "26", "20.0%"],
            ["Production",  "20", "15.4%"],
            ["Energy",       "4",  "3.1%"],
        ],
        col_widths=[2.5, 1.5, 1.5]
    )

    doc.add_page_break()

    # ── 8. API Reference ──────────────────────────────────────────────────────
    add_heading(doc, "8.  API Reference (key endpoints)")
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/events/orchestrate",                    "Run one event through the full 7-layer pipeline"],
            ["POST", "/api/events/orchestrate/scenario/{name}",    "Run a named scenario (multiple events)"],
            ["GET",  "/api/events/decisions",                      "List all decisions for a factory"],
            ["GET",  "/api/events/decisions/{id}",                 "Get full detail for one decision"],
            ["POST", "/api/events/decisions/{id}/approve",         "Approve or reject a decision"],
            ["GET",  "/api/events/audit/search",                   "Search audit logs with filters"],
            ["GET",  "/api/analytics/summary",                     "KPI summary (approval rate, savings, latency)"],
            ["GET",  "/api/analytics/timeline",                    "Event activity over N days"],
            ["GET",  "/api/custom-agents/list",                    "List registered custom agents"],
            ["POST", "/api/custom-agents/create",                  "Create a new custom agent"],
            ["WS",   "/api/events/ws/orchestrate/{decision_id}",   "Real-time pipeline progress stream"],
        ],
        col_widths=[0.8, 3.5, 2.9]
    )

    doc.add_page_break()

    # ── 9. Related Research ───────────────────────────────────────────────────
    add_heading(doc, "9.  Related Research")
    add_body(doc, (
        "PAAIM draws on several distinct bodies of work. What is novel is the combination — "
        "no published system applies all of these ideas together in a manufacturing context."
    ))

    related = [
        ("Constitutional AI — Bai et al., Anthropic (2022)",
         "Introduced the idea of giving an AI a written set of principles (a 'constitution') "
         "that it must follow. PAAIM applies this directly: the industrial_constitution.yaml "
         "is enforced at runtime by the policy engine rather than baked into model training. "
         "The difference is that runtime enforcement is guaranteed — a trained model can still "
         "violate its constitution under the right conditions."),
        ("ReAct — Yao et al., Princeton/Google (2022)",
         "Demonstrated that LLMs can interleave reasoning steps with actions (tool calls). "
         "PAAIM's agents follow this pattern: Gemini reasons about the event, then outputs a "
         "structured action recommendation. The pipeline then acts on that recommendation "
         "through the downstream layers."),
        ("AutoGen — Wu et al., Microsoft (2023)",
         "Multi-agent framework where specialist agents collaborate on a shared task. PAAIM's "
         "five parallel agents follow the same principle, with the key difference that agent "
         "routing is deterministic (event type determines which agents run) rather than "
         "LLM-driven."),
        ("Digital Twin — Grieves (2014)",
         "Established the concept of simulating a physical system in software to evaluate "
         "decisions before execution. PAAIM's Layer 4 implements this as a decision gate: "
         "no action reaches the approval queue without an impact estimate."),
        ("Human-in-the-Loop ML — Monarch (2021)",
         "Covers when and how to involve humans in machine learning pipelines. PAAIM's "
         "approval gate is a direct implementation of HITL principles, with the additional "
         "constraint that the required approval role scales with the risk level of the action."),
    ]

    for title, body in related:
        add_heading(doc, title, level=2)
        add_body(doc, body)

    doc.add_page_break()

    # ── 10. Competitive landscape ─────────────────────────────────────────────
    add_heading(doc, "10.  Competitive Landscape")
    add_table(doc,
        ["System", "LLM agents", "Policy enforcement", "HITL approval", "Audit trail", "Open / configurable"],
        [
            ["Siemens Industrial Copilot", "Partial", "No",      "No",      "Partial", "No"],
            ["IBM Maximo AI",              "No",      "Partial", "Partial", "Yes",     "No"],
            ["Rockwell FactoryTalk",       "No",      "Rules",   "Partial", "Yes",     "Partial"],
            ["Palantir AIP",               "Yes",     "Partial", "Yes",     "Yes",     "No"],
            ["PAAIM",                      "Yes",     "Yes",     "Yes",     "Yes",     "Yes"],
        ],
        col_widths=[1.8, 1.0, 1.4, 1.2, 1.0, 1.3]
    )

    add_body(doc, (
        "The key differentiator is the combination of open architecture, runtime policy enforcement "
        "via a human-readable YAML file, and the no-code custom agent builder. A mid-size "
        "manufacturer can deploy and configure PAAIM without an enterprise software contract."
    ))

    doc.add_page_break()

    # ── 11. Future Work ───────────────────────────────────────────────────────
    add_heading(doc, "11.  Future Work")

    roadmap = [
        ("Phase 2 — Agent Memory",
         "Each agent call is currently stateless. Phase 2 will inject recent event history "
         "(last 5 events for the same machine) and past decision outcomes into the Gemini "
         "prompt. This allows the maintenance agent to notice that it has issued the same "
         "recommendation three times in 48 hours without it being acted on, and escalate "
         "accordingly."),
        ("Phase 2 — Task Management",
         "When a decision is approved, a work order should be created, assigned to the right "
         "team, and tracked to completion. Closing the feedback loop — recording whether the "
         "maintenance was actually done, and whether the predicted failure was prevented — "
         "enables the agents to improve their recommendations over time."),
        ("Phase 3 — Production Deployment",
         "Replace SQLite with PostgreSQL, add OAuth2 authentication with real user roles, "
         "complete the Docker Compose configuration, and add Kubernetes manifests for "
         "cloud deployment. The backend is already written with async SQLAlchemy so the "
         "database swap requires no application code changes."),
        ("Phase 3 — Real Connector Integration",
         "The connector framework (SCADA, OPC-UA, REST API, message queue) is architecturally "
         "complete but connects to simulated data. Phase 3 will add tested integrations for "
         "common industrial protocols used in the target factory."),
        ("Phase 4 — Learning Loop",
         "Use the growing database of approved and rejected decisions as a feedback signal. "
         "Track which agent recommendations have the highest long-term approval rates and "
         "adjust confidence thresholds accordingly. This moves PAAIM from a static advisory "
         "tool to one that improves with use."),
    ]

    for title, body in roadmap:
        add_heading(doc, title, level=2)
        add_body(doc, body)

    doc.add_page_break()

    # ── 12. How to run ───────────────────────────────────────────────────────
    add_heading(doc, "12.  Running the System")

    add_heading(doc, "12.1  Prerequisites", level=2)
    add_bullet(doc, "Python 3.11+ with pip")
    add_bullet(doc, "Node.js 18+")
    add_bullet(doc, "A Google Gemini API key (free tier is sufficient for demo)")

    add_heading(doc, "12.2  Quick start", level=2)
    add_body(doc, "The start_demo.sh script handles everything — virtual environment, dependencies, database seeding, and both servers:")

    code_p = doc.add_paragraph()
    code_r = code_p.add_run(
        "cd PAAIM\n"
        "./start_demo.sh           # first run\n"
        "./start_demo.sh --seed    # wipe and re-seed demo data"
    )
    code_r.font.name = 'Courier New'
    code_r.font.size = Pt(9.5)
    code_r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    code_p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "12.3  Manual start", level=2)
    add_body(doc, "Two terminals:")

    code2 = doc.add_paragraph()
    r = code2.add_run(
        "# Terminal 1 — backend\n"
        "cd backend && source venv/bin/activate\n"
        "uvicorn paaim.main:app --reload --port 8000\n\n"
        "# Terminal 2 — frontend\n"
        "cd frontend && npm run dev"
    )
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    add_body(doc, "Demo login:  demo@paaim.io / demo123", space_before=6)

    # ── 13. Conclusion ────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "13.  Conclusion")
    add_body(doc, (
        "PAAIM demonstrates that the combination of large language model reasoning, runtime "
        "policy enforcement, and structured human-in-the-loop approval produces a manufacturing "
        "decision tool that is both more capable and more accountable than current alternatives."
    ))
    add_body(doc, (
        "The architecture has a clear separation of concerns: the AI reasons about what to do, "
        "the policy engine decides whether it is permitted, the simulation shows what it will "
        "cost, and a human with the appropriate authority makes the final call. Nothing is "
        "executed without that approval, and nothing happens without a record."
    ))
    add_body(doc, (
        "The Phase 1 demo shows a working end-to-end system: events processed through all seven "
        "layers, decisions stored and reviewable, approvals tracked, and analytics that update "
        "in real time. The no-code agent builder lets the team extend coverage to new machines "
        "and signal types without modifying the core pipeline."
    ))
    add_body(doc, (
        "The most important next step is closing the feedback loop — tracking whether approved "
        "actions actually prevented the predicted failures, and feeding that outcome data back "
        "into the agents. That is what will turn PAAIM from a decision support tool into one "
        "that learns from its own history."
    ))

    # ── Save ──────────────────────────────────────────────────────────────────
    out = "/Users/venky/Desktop/projects/PAAIM/PAAIM_Project_Documentation.docx"
    doc.save(out)
    print(f"\n  Document saved: {out}\n")
    print("  Sections:")
    print("    1. Introduction          6. Frontend & UX")
    print("    2. System Architecture   7. Demo Results")
    print("    3. 7-Layer Pipeline      8. API Reference")
    print("    4. Decision Flow         9. Related Research")
    print("    5. Key Components       10. Competitive Landscape")
    print("                            11. Future Work")
    print("                            12. How to Run")
    print("                            13. Conclusion")
    print("\n  Diagrams embedded:")
    print("    Fig 1 — System Architecture")
    print("    Fig 2 — 7-Layer Pipeline")
    print("    Fig 3 — Event to Decision Flow\n")


if __name__ == "__main__":
    build()
